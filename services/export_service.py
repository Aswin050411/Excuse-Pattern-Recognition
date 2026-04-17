import csv
import io
import os
from flask import make_response, current_app
from repository.db import execute_query
from flask import make_response, current_app
from repository.db import execute_query
from services.analytics_service import get_analytics_data

def generate_csv_report(user_id=None, role='employee', type='tasks'):
    """Generates a CSV report with Excel-compatible UTF-8 BOM."""
    headers, data = _fetch_report_data(user_id, role, type)
    
    si = io.StringIO()
    si.write('\ufeff') # UTF-8 BOM
    cw = csv.writer(si)
    cw.writerow(headers)
    for row in data:
        cw.writerow(list(row.values()) if isinstance(row, dict) else list(row))
        
    response = make_response(si.getvalue())
    response.headers["Content-Disposition"] = f"attachment; filename={type}_report.csv"
    response.headers["Content-type"] = "text/csv; charset=utf-8-sig"
    return response

def generate_word_report(user_id=None, role='employee', type='tasks'):
    """Generates a Word (.docx) report."""
    from docx import Document
    headers, data = _fetch_report_data(user_id, role, type)
    
    doc = Document()
    doc.add_heading(f'Excuse Pattern AI - {type.capitalize()} Report', 0)
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        
    for row in data:
        row_cells = table.add_row().cells
        row_values = list(row.values()) if isinstance(row, dict) else list(row)
        for i, val in enumerate(row_values):
            row_cells[i].text = str(val)
            
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    
    response = make_response(target_stream.getvalue())
    response.headers["Content-Disposition"] = f"attachment; filename={type}_report.docx"
    response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return response

def generate_pdf_report(user_id=None, role='employee', type='tasks'):
    """Generates a PDF report."""
    from fpdf import FPDF
    headers, data = _fetch_report_data(user_id, role, type)
    
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, f'Excuse Pattern AI - {type.capitalize()} Report', ln=True, align='C')
    pdf.ln(10)
    
    # Table headers
    pdf.set_font("Helvetica", 'B', 10)
    col_width = pdf.epw / len(headers)
    for header in headers:
        pdf.cell(col_width, 10, header, border=1)
    pdf.ln()
    
    # Table data
    pdf.set_font("Helvetica", size=8)
    for row in data:
        row_values = list(row.values()) if isinstance(row, dict) else list(row)
        for val in row_values:
            pdf.cell(col_width, 8, str(val)[:30], border=1) # Truncate for A4 width
        pdf.ln()
        
    response = make_response(pdf.output())
    response.headers["Content-Disposition"] = f"attachment; filename={type}_report.pdf"
    response.headers["Content-type"] = "application/pdf"
    return response

def generate_analytics_pdf_report(user_id=None, role='employee'):
    """Generates a comprehensive PDF report based on all analytics data."""
    from fpdf import FPDF
    kpis = get_analytics_data(user_id=user_id, role=role)
    
    def safe_text(text):
        if not text: return ""
        return str(text).encode('latin-1', 'replace').decode('latin-1')

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    
    # Header
    pdf.set_fill_color(30, 41, 59) # Slate 800
    pdf.rect(0, 0, 210, 40, 'F')
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", 'B', 24)
    pdf.cell(0, 20, safe_text('EXCUSE PATTERN AI'), ln=True, align='C')
    pdf.set_font("Helvetica", '', 12)
    pdf.cell(0, 10, safe_text('Comprehensive Analytics Intelligence Report'), ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", 'B', 14)
    pdf.cell(0, 10, safe_text(f"Scope: {'Team' if role in ('admin', 'manager') else 'Personal'} Analytics"), ln=True)
    pdf.set_font("Helvetica", '', 10)
    pdf.cell(0, 5, safe_text(f"Generated for User ID: {user_id}"), ln=True)
    pdf.ln(5)

    # 1. CORE KPIs
    _pdf_section_header(pdf, "CORE RELIABILITY METRICS")
    cols = [
        ("Reliability (WRS)", f"{kpis.get('wrs_score', 0)}"),
        ("Delay Rate", f"{kpis.get('delay_rate', 0)}%"),
        ("Total Delays", f"{kpis.get('total_delays', 0)}"),
        ("Avg Authenticity", f"{kpis.get('avg_auth_score', 0)}%")
    ]
    for label, val in cols:
        pdf.set_font("Helvetica", 'B', 10)
        pdf.cell(45, 8, safe_text(label), border=1)
        pdf.set_font("Helvetica", '', 10)
        pdf.cell(45, 8, safe_text(val), border=1, ln=True)
    pdf.ln(10)

    # 2. AI INTELLIGENCE
    _pdf_section_header(pdf, "AI INTELLIGENCE & HEURISTICS")
    
    # Executive Summary
    pdf.set_font("Helvetica", 'B', 11)
    pdf.cell(0, 8, safe_text("Executive Summary:"), ln=True)
    pdf.set_font("Helvetica", 'I', 10)
    pdf.multi_cell(0, 6, safe_text(kpis.get('executive_summary', "No summary available.")))
    pdf.ln(5)
    
    # Advanced Intelligence
    ai = kpis.get('ai', {})
    intel_cols = [
        ("Behavioral Intel Score", f"{ai.get('intel_score', 0)}"),
        ("Risk Momentum", f"{ai.get('momentum_ai', {}).get('risk_momentum', 'Stable')}"),
        ("Delay Probability", f"{ai.get('prediction_ai', {}).get('delay_probability', 0)}%"),
        ("Excuse Quality", f"{ai.get('quality_ai', {}).get('quality_label', 'Neutral')}")
    ]
    for label, val in intel_cols:
        pdf.set_font("Helvetica", 'B', 10)
        pdf.cell(50, 8, safe_text(label), border=1)
        pdf.set_font("Helvetica", '', 10)
        pdf.cell(40, 8, safe_text(val), border=1, ln=True)
    pdf.ln(5)

    # 3. AI INSIGHTS
    _pdf_section_header(pdf, "STRATEGIC INSIGHTS")
    for insight in kpis.get('ai_insights', []):
        severity = insight.get('severity', 'info').upper()
        pdf.set_font("Helvetica", 'B', 9)
        pdf.cell(30, 6, safe_text(f"[{severity}]"), ln=0)
        pdf.set_font("Helvetica", '', 9)
        pdf.multi_cell(0, 6, safe_text(insight.get('text', "")))
    
    pdf.ln(10)

    # Footer
    pdf.set_y(-15)
    pdf.set_font("Helvetica", 'I', 8)
    pdf.cell(0, 10, safe_text('Proprietary AI Analysis - Experimental Demo Only'), align='C')

    response = make_response(pdf.output())
    response.headers["Content-Disposition"] = "attachment; filename=comprehensive_analytics.pdf"
    response.headers["Content-type"] = "application/pdf"
    return response

def _pdf_section_header(pdf, title):
    pdf.set_fill_color(226, 232, 240) # Slate 200
    pdf.set_font("Helvetica", 'B', 12)
    pdf.cell(0, 10, f" {title}", ln=True, fill=True)
    pdf.ln(2)

def _fetch_report_data(user_id, role, type):
    """Internal helper to fetch data for all formats."""
    data = []
    headers = []
    try:
        if type == 'tasks':
            where_clause = ""
            params = []
            if role == 'employee' and user_id:
                where_clause = "WHERE assigned_to = %s"
                params = [user_id]
            data = execute_query(f"SELECT id, title, status, priority, deadline FROM tasks {where_clause} ORDER BY created_at DESC", params)
            headers = ['ID', 'Title', 'Status', 'Priority', 'Deadline']
        else:
            where_clause = ""
            params = []
            if role == 'employee' and user_id:
                where_clause = "WHERE d.user_id = %s"
                params = [user_id]
            data = execute_query(f"""
                SELECT d.id, t.title, d.score_authenticity, d.risk_level, d.submitted_at 
                FROM delays d LEFT JOIN tasks t ON d.task_id = t.id 
                {where_clause} ORDER BY d.submitted_at DESC
            """, params)
            headers = ['ID', 'Task', 'Auth Score', 'Risk', 'Date']
    except Exception as e:
        current_app.logger.error(f"Export data fetch error: {e}")
    
    return headers, (data if data else [])
